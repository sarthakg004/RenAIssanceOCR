"""
Gemini OCR Backend Server
FastAPI server for OCR text recognition using Google Gemini
"""

import os
import io
import time
import base64
from typing import Optional
from datetime import datetime

from fastapi import FastAPI, File, UploadFile, HTTPException, Header, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from starlette.datastructures import UploadFile as StarletteUploadFile
from starlette.requests import Request as StarletteRequest

# Increase max upload size to 100MB (Gemini supports up to 100MB images)
MAX_UPLOAD_SIZE = 100 * 1024 * 1024  # 100MB in bytes

import cv2
import numpy as np
import json

from google import genai
from google.genai import types

from docx import Document

# Import preprocessing module
from preprocessing import run_pipeline, OP_REGISTRY, validate_pipeline_config
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch

app = FastAPI(title="Gemini OCR API", version="1.0.0")

# CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:5174", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================
# Rate Limiting - Simple In-Memory Guard
# ============================================

class RateLimiter:
    """Simple rate limiter: 5 requests per minute = 1 request per 12 seconds"""
    
    def __init__(self, min_interval: int = 12):
        self.min_interval = min_interval
        self.last_request_time: Optional[float] = None
    
    def can_proceed(self) -> tuple[bool, int]:
        """Check if request can proceed. Returns (can_proceed, wait_time_seconds)"""
        if self.last_request_time is None:
            return True, 0
        
        elapsed = time.time() - self.last_request_time
        if elapsed >= self.min_interval:
            return True, 0
        
        wait_time = int(self.min_interval - elapsed) + 1
        return False, wait_time
    
    def record_request(self):
        """Record that a request was made"""
        self.last_request_time = time.time()
    
    def get_status(self) -> dict:
        """Get current rate limit status"""
        if self.last_request_time is None:
            return {"ready": True, "wait_seconds": 0, "last_request": None}
        
        elapsed = time.time() - self.last_request_time
        if elapsed >= self.min_interval:
            return {"ready": True, "wait_seconds": 0, "last_request": self.last_request_time}
        
        wait_time = int(self.min_interval - elapsed) + 1
        return {"ready": False, "wait_seconds": wait_time, "last_request": self.last_request_time}

# Global rate limiter instance - 5 requests per minute = 12 second intervals
rate_limiter = RateLimiter(min_interval=12)

# ============================================
# Gemini OCR Core
# ============================================

AVAILABLE_MODELS = [
    {
        "id": "gemini-3-flash-preview",
        "name": "Gemini 3 Flash Preview",
        "description": "Latest and fastest preview model (recommended)"
    },
    {
        "id": "gemini-3-pro-preview",
        "name": "Gemini 3 Pro Preview",
        "description": "Most capable preview model for complex documents"
    },
    {
        "id": "gemini-2.5-pro",
        "name": "Gemini 2.5 Pro",
        "description": "Stable pro model, excellent accuracy"
    },
    {
        "id": "gemini-2.5-flash", 
        "name": "Gemini 2.5 Flash",
        "description": "Stable flash model, good balance of speed and quality"
    },
]

DEFAULT_MODEL = "gemini-3-flash-preview"

# List of model IDs for validation
MODEL_IDS = [m["id"] for m in AVAILABLE_MODELS]

def get_gemini_client(api_key: str):
    """Create Gemini client with provided API key"""
    return genai.Client(api_key=api_key)

def gemini_ocr(client, image_bytes: bytes, model_name: str, mime_type: str = "image/png") -> str:
    prompt = """
        You are performing high-accuracy OCR transcription.

        Transcribe ALL readable text exactly as it appears in the image.

        Core rules:

        * Preserve original line breaks.
        * Preserve paragraph spacing.
        * Preserve punctuation and special characters.
        * Preserve original spelling (do NOT modernize).
        * Preserve capitalization exactly.
        * Keep hyphenated line-break words exactly as shown.
        * Do NOT summarize.
        * Do NOT explain.
        * Output only the transcription.

        Layout rules:

        * If text is in multiple columns, transcribe column by column from left to right.
        * Preserve indentation if visible.
        * Keep headings and section breaks.
        * Keep marginal notes or side text on separate lines and prefix them with "[margin] ".

        Context-based reconstruction rules:

        * If a word is partially unclear, use surrounding letters and sentence context to infer the most likely word.
        * Prefer historically and linguistically plausible words over random guesses.
        * Use your language knowledge to reconstruct faded or broken characters when confidence is reasonably high.
        * Do NOT mark a word as illegible if a strong contextual reconstruction is possible.

        Uncertainty handling:

        * If reconstruction is reasonably confident → output the reconstructed word normally.
        * If multiple interpretations are possible → choose the most contextually likely one.
        * If text is truly unreadable with no strong contextual clue → use [illegible].
        * If only one or two characters are unclear but the word is inferable → output the full inferred word.

        Noise handling:

        * Ignore page borders, stains, ornaments, and decorative lines.
        * Do not include printer marks unless they are clearly text.
    """

    response = client.models.generate_content(
        model=model_name,
        contents=[
            types.Part.from_bytes(data=image_bytes, mime_type=mime_type),
            prompt
        ],
    )

    return response.text


# ============================================
# API Endpoints
# ============================================

class OCRResponse(BaseModel):
    success: bool
    transcript: Optional[str] = None
    error: Optional[str] = None
    model_used: str
    processing_time_ms: int

class RateLimitResponse(BaseModel):
    ready: bool
    wait_seconds: int
    last_request: Optional[float]

class OCRRequest(BaseModel):
    """Request model for JSON-based OCR endpoint (supports large images)"""
    image_data: str  # Base64 encoded image with data URL prefix
    model: str = DEFAULT_MODEL

class ExportRequest(BaseModel):
    transcripts: dict  # {page_number: transcript_text}
    format: str  # "txt", "docx", "pdf"


class PreprocessRequest(BaseModel):
    """Request model for preprocessing endpoint"""
    image_data: str  # Base64 encoded image with data URL prefix
    operations: list  # List of {op, params, enabled} dicts
    preview_mode: bool = False  # Use faster algorithms for preview


# ============================================
# Preprocessing Endpoints
# ============================================

@app.get("/api/preprocess/operations")
async def get_available_operations():
    """Get list of available preprocessing operations"""
    return {
        "operations": list(OP_REGISTRY.keys()),
        "descriptions": {
            "normalize": "Normalize image brightness and contrast levels",
            "grayscale": "Convert image to grayscale",
            "deskew": "Automatically correct image rotation/skew",
            "denoise": "Remove noise while preserving text edges",
            "contrast": "Enhance contrast using CLAHE",
            "sharpen": "Sharpen text edges for clearer text",
            "threshold": "Convert to binary (black and white)",
        }
    }


@app.post("/api/preprocess")
async def preprocess_image_endpoint(request: PreprocessRequest):
    """
    Apply preprocessing pipeline to an image.
    
    JSON Body:
        image_data: Base64 encoded image (with or without data URL prefix)
        operations: List of operations to apply
        preview_mode: Use faster algorithms for preview (optional)
    
    Returns:
        Processed image as base64, processing info
    """
    start_time = time.time()
    
    try:
        # Validate operations config
        validation = validate_pipeline_config(request.operations)
        if not validation["valid"]:
            raise HTTPException(
                status_code=400,
                detail={
                    "error": "invalid_operations",
                    "message": "Invalid pipeline configuration",
                    "errors": validation["errors"]
                }
            )
        
        # Parse base64 image data
        image_data = request.image_data
        if "," in image_data:
            # Has data URL prefix like "data:image/png;base64,..."
            header, encoded = image_data.split(",", 1)
            mime_type = header.split(":")[1].split(";")[0] if ":" in header else "image/png"
        else:
            encoded = image_data
            mime_type = "image/png"
        
        # Decode image
        image_bytes = base64.b64decode(encoded)
        nparr = np.frombuffer(image_bytes, np.uint8)
        image = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if image is None:
            raise HTTPException(
                status_code=400,
                detail={"error": "invalid_image", "message": "Could not decode image"}
            )
        
        # Run preprocessing pipeline
        result = run_pipeline(
            image=image,
            steps=request.operations,
            continue_on_error=True,
            preview_mode=request.preview_mode,
        )
        
        # Encode result image
        if result.image is not None:
            # Determine output format based on input
            if "jpeg" in mime_type or "jpg" in mime_type:
                encode_param = [cv2.IMWRITE_JPEG_QUALITY, 95]
                _, buffer = cv2.imencode('.jpg', result.image, encode_param)
                output_mime = "image/jpeg"
            else:
                _, buffer = cv2.imencode('.png', result.image)
                output_mime = "image/png"
            
            encoded_result = base64.b64encode(buffer).decode('utf-8')
            result_data_url = f"data:{output_mime};base64,{encoded_result}"
        else:
            result_data_url = None
        
        processing_time = int((time.time() - start_time) * 1000)
        
        return {
            "success": result.success,
            "processed_image": result_data_url,
            "processing_time_ms": processing_time,
            "progress_info": result.progress_info,
            "errors": [
                {"step": e["step"], "error": e["error"]}
                for e in result.errors
            ] if result.errors else [],
        }
        
    except HTTPException:
        raise
    except Exception as e:
        processing_time = int((time.time() - start_time) * 1000)
        return {
            "success": False,
            "processed_image": None,
            "processing_time_ms": processing_time,
            "error": str(e),
        }


@app.post("/api/preprocess/validate")
async def validate_operations(operations: list):
    """
    Validate preprocessing pipeline configuration.
    
    JSON Body:
        List of operations to validate
    
    Returns:
        Validation result with any errors
    """
    validation = validate_pipeline_config(operations)
    return validation

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

@app.get("/api/models")
async def get_available_models():
    """Get list of available Gemini models"""
    return {
        "models": AVAILABLE_MODELS,
        "default": DEFAULT_MODEL
    }

@app.post("/api/validate-key")
async def validate_api_key(x_gemini_api_key: str = Header(..., alias="X-Gemini-API-Key")):
    """
    Validate a Gemini API key format.
    
    Note: We do NOT make an API call to validate the key because:
    1. It wastes quota (rate limits)
    2. The key will be validated on first actual OCR request
    3. Rate limit errors would falsely mark valid keys as invalid
    
    We only check the format here. The actual validation happens
    when processing pages - if the key is invalid, that call will fail.
    """
    if not x_gemini_api_key:
        raise HTTPException(status_code=401, detail="API key is required")
    
    # Gemini API keys typically start with 'AI' and are ~39 characters
    # But we'll be lenient and just check minimum length
    if len(x_gemini_api_key) < 20:
        raise HTTPException(status_code=401, detail="API key appears too short")
    
    # Format looks valid - actual validation will happen on first OCR call
    return {
        "valid": True, 
        "message": "API key format is valid. It will be verified on first use."
    }

@app.get("/api/rate-limit-status")
async def get_rate_limit_status():
    """Get current rate limit status"""
    return rate_limiter.get_status()

@app.post("/api/gemini-ocr-page", response_model=OCRResponse)
async def ocr_page(
    image: UploadFile = File(...),
    model: str = Form(default=DEFAULT_MODEL),
    x_gemini_api_key: str = Header(..., alias="X-Gemini-API-Key")
):
    """
    Process a single page image with Gemini OCR
    
    Headers:
        X-Gemini-API-Key: Your Gemini API key
    
    Form Data:
        image: Image file to process
        model: Gemini model name (optional, defaults to gemini-2.0-flash)
    """
    # Check rate limit
    can_proceed, wait_time = rate_limiter.can_proceed()
    if not can_proceed:
        raise HTTPException(
            status_code=429,
            detail={
                "error": "rate_limited",
                "message": f"Rate limit exceeded. Please wait {wait_time} seconds.",
                "wait_seconds": wait_time
            }
        )
    
    # Validate model
    if model not in MODEL_IDS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid model. Available models: {MODEL_IDS}"
        )
    
    # Validate API key
    if not x_gemini_api_key or len(x_gemini_api_key) < 10:
        raise HTTPException(
            status_code=401,
            detail="Invalid or missing Gemini API key"
        )
    
    start_time = time.time()
    
    try:
        # Read image
        image_bytes = await image.read()
        
        # Determine MIME type
        content_type = image.content_type or "image/png"
        if content_type not in ["image/png", "image/jpeg", "image/jpg", "image/webp"]:
            content_type = "image/png"
        
        # Create client and perform OCR
        client = get_gemini_client(x_gemini_api_key)
        transcript = gemini_ocr(client, image_bytes, model, content_type)
        
        # Record successful request for rate limiting
        rate_limiter.record_request()
        
        processing_time = int((time.time() - start_time) * 1000)
        
        return OCRResponse(
            success=True,
            transcript=transcript,
            model_used=model,
            processing_time_ms=processing_time
        )
        
    except Exception as e:
        error_msg = str(e)
        
        # Check for specific API errors
        if "API_KEY_INVALID" in error_msg or "401" in error_msg:
            raise HTTPException(status_code=401, detail="Invalid Gemini API key")
        if "QUOTA_EXCEEDED" in error_msg or "429" in error_msg:
            raise HTTPException(status_code=429, detail="Gemini API quota exceeded")
        
        processing_time = int((time.time() - start_time) * 1000)
        
        return OCRResponse(
            success=False,
            error=error_msg,
            model_used=model,
            processing_time_ms=processing_time
        )

@app.post("/api/gemini-ocr-base64", response_model=OCRResponse)
async def ocr_page_base64(
    image_data: str = Form(...),
    model: str = Form(default=DEFAULT_MODEL),
    x_gemini_api_key: str = Header(..., alias="X-Gemini-API-Key")
):
    """
    Process a base64-encoded image with Gemini OCR
    
    Headers:
        X-Gemini-API-Key: Your Gemini API key
    
    Form Data:
        image_data: Base64 encoded image (with or without data URL prefix)
        model: Gemini model name
    """
    # Check rate limit
    can_proceed, wait_time = rate_limiter.can_proceed()
    if not can_proceed:
        raise HTTPException(
            status_code=429,
            detail={
                "error": "rate_limited",
                "message": f"Rate limit exceeded. Please wait {wait_time} seconds.",
                "wait_seconds": wait_time
            }
        )
    
    # Validate model
    if model not in MODEL_IDS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid model. Available models: {MODEL_IDS}"
        )
    
    start_time = time.time()
    
    try:
        # Parse base64 data
        if "," in image_data:
            # Has data URL prefix like "data:image/png;base64,..."
            header, encoded = image_data.split(",", 1)
            mime_type = header.split(":")[1].split(";")[0] if ":" in header else "image/png"
        else:
            encoded = image_data
            mime_type = "image/png"
        
        image_bytes = base64.b64decode(encoded)
        
        # Create client and perform OCR
        client = get_gemini_client(x_gemini_api_key)
        transcript = gemini_ocr(client, image_bytes, model, mime_type)
        
        # Record successful request
        rate_limiter.record_request()
        
        processing_time = int((time.time() - start_time) * 1000)
        
        return OCRResponse(
            success=True,
            transcript=transcript,
            model_used=model,
            processing_time_ms=processing_time
        )
        
    except Exception as e:
        error_msg = str(e)
        processing_time = int((time.time() - start_time) * 1000)
        
        if "API_KEY_INVALID" in error_msg or "401" in error_msg:
            raise HTTPException(status_code=401, detail="Invalid Gemini API key")
        
        return OCRResponse(
            success=False,
            error=error_msg,
            model_used=model,
            processing_time_ms=processing_time
        )


@app.post("/api/gemini-ocr-json", response_model=OCRResponse)
async def ocr_page_json(
    request: OCRRequest,
    x_gemini_api_key: str = Header(..., alias="X-Gemini-API-Key")
):
    """
    Process a base64-encoded image with Gemini OCR (JSON body)
    
    This endpoint uses JSON body instead of FormData, which allows for 
    larger image uploads (up to 100MB) without hitting multipart size limits.
    
    Headers:
        X-Gemini-API-Key: Your Gemini API key
    
    JSON Body:
        image_data: Base64 encoded image (with or without data URL prefix)
        model: Gemini model name
    """
    # Check rate limit
    can_proceed, wait_time = rate_limiter.can_proceed()
    if not can_proceed:
        raise HTTPException(
            status_code=429,
            detail={
                "error": "rate_limited",
                "message": f"Rate limit exceeded. Please wait {wait_time} seconds.",
                "wait_seconds": wait_time
            }
        )
    
    # Validate model
    if request.model not in MODEL_IDS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid model. Available models: {MODEL_IDS}"
        )
    
    start_time = time.time()
    
    try:
        image_data = request.image_data
        
        # Parse base64 data
        if "," in image_data:
            # Has data URL prefix like "data:image/png;base64,..."
            header, encoded = image_data.split(",", 1)
            mime_type = header.split(":")[1].split(";")[0] if ":" in header else "image/png"
        else:
            encoded = image_data
            mime_type = "image/png"
        
        image_bytes = base64.b64decode(encoded)
        
        # Create client and perform OCR
        client = get_gemini_client(x_gemini_api_key)
        transcript = gemini_ocr(client, image_bytes, request.model, mime_type)
        
        # Record successful request
        rate_limiter.record_request()
        
        processing_time = int((time.time() - start_time) * 1000)
        
        return OCRResponse(
            success=True,
            transcript=transcript,
            model_used=request.model,
            processing_time_ms=processing_time
        )
        
    except Exception as e:
        error_msg = str(e)
        processing_time = int((time.time() - start_time) * 1000)
        
        if "API_KEY_INVALID" in error_msg or "401" in error_msg:
            raise HTTPException(status_code=401, detail="Invalid Gemini API key")
        
        return OCRResponse(
            success=False,
            error=error_msg,
            model_used=request.model,
            processing_time_ms=processing_time
        )


# ============================================
# Export Endpoints
# ============================================

def build_combined_transcript(transcripts: dict) -> str:
    """Build combined transcript with page separators"""
    pages = sorted(transcripts.keys(), key=lambda x: int(x) if x.isdigit() else 0)
    
    sections = []
    for page in pages:
        text = transcripts[page]
        section = f"page_{page}\n{'─' * 20}\n{text}"
        sections.append(section)
    
    return "\n\n".join(sections)

@app.post("/api/export/txt")
async def export_txt(request: ExportRequest):
    """Export combined transcript as TXT file"""
    combined = build_combined_transcript(request.transcripts)
    
    buffer = io.BytesIO(combined.encode('utf-8'))
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="text/plain",
        headers={"Content-Disposition": "attachment; filename=transcript_full.txt"}
    )

@app.post("/api/export/docx")
async def export_docx(request: ExportRequest):
    """Export combined transcript as DOCX file"""
    doc = Document()
    
    # Set document title
    title = doc.add_heading("Combined Transcript", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacer
    
    # Add each page
    pages = sorted(request.transcripts.keys(), key=lambda x: int(x) if x.isdigit() else 0)
    
    for i, page in enumerate(pages):
        text = request.transcripts[page]
        
        # Page header
        heading = doc.add_heading(f"Page {page}", level=1)
        
        # Separator line
        separator = doc.add_paragraph("─" * 40)
        separator.runs[0].font.size = Pt(10)
        
        # Page content
        for para_text in text.split('\n'):
            if para_text.strip():
                para = doc.add_paragraph(para_text)
                para.style.font.size = Pt(11)
        
        # Add page break between pages (except last)
        if i < len(pages) - 1:
            doc.add_page_break()
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=transcript_full.docx"}
    )

@app.post("/api/export/pdf")
async def export_pdf(request: ExportRequest):
    """Export combined transcript as PDF file"""
    buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center
    )
    
    page_header_style = ParagraphStyle(
        'PageHeader',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=6,
        textColor='#1e40af'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6,
        leading=14
    )
    
    separator_style = ParagraphStyle(
        'Separator',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=12,
        textColor='#6b7280'
    )
    
    story = []
    
    # Title
    story.append(Paragraph("Combined Transcript", title_style))
    story.append(Spacer(1, 20))
    
    # Add each page
    pages = sorted(request.transcripts.keys(), key=lambda x: int(x) if x.isdigit() else 0)
    
    for page in pages:
        text = request.transcripts[page]
        
        # Page header
        story.append(Paragraph(f"Page {page}", page_header_style))
        story.append(Paragraph("─" * 50, separator_style))
        
        # Page content - escape HTML entities
        for line in text.split('\n'):
            if line.strip():
                # Escape special characters for reportlab
                safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_line, body_style))
            else:
                story.append(Spacer(1, 6))
        
        story.append(Spacer(1, 30))
    
    doc.build(story)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=transcript_full.pdf"}
    )

# ============================================
# Run Server
# ============================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
